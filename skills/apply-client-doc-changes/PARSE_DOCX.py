#!/usr/bin/env python3
"""
apply-client-doc-changes · 解析客戶 docx 改稿
用法：
  python3 PARSE_DOCX.py <docx_path>
  python3 PARSE_DOCX.py /tmp/client-doc-1qfsAe.docx

輸出：
  ★ 黃字段落（新增/修改）
  ✂ 刪除線段落（要砍）
  📷 嵌入圖片清單
"""
import sys
import zipfile
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: pip install python-docx")
    sys.exit(1)


def analyze_paragraph(p):
    """回傳該段的 yellow 文字 + strike 文字"""
    yellow = "".join(
        r.text for r in p.runs
        if r.font.highlight_color and "YELLOW" in str(r.font.highlight_color)
    )
    strike = "".join(r.text for r in p.runs if r.font.strike)
    return yellow, strike


def walk_tables(doc):
    """表格內也可能有黃字 / 刪除線"""
    out = []
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    y, s = analyze_paragraph(p)
                    if y.strip() or s.strip():
                        out.append((f"T{ti}.R{ri}.C{ci}.P{pi}", p.text, y, s))
    return out


def main(docx_path: str):
    p = Path(docx_path)
    if not p.exists():
        print(f"ERROR: {docx_path} not found")
        sys.exit(1)

    d = Document(str(p))

    # 段落
    yellow_paras = []
    strike_paras = []
    for i, para in enumerate(d.paragraphs):
        y, s = analyze_paragraph(para)
        if y.strip():
            yellow_paras.append((i, para.text, y))
        if s.strip():
            strike_paras.append((i, para.text, s))

    # 表格
    tbl = walk_tables(d)

    # 嵌入圖片
    with zipfile.ZipFile(str(p)) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/")]

    # 輸出
    print(f"=== 黃字段落（新增/修改）共 {len(yellow_paras)} 處 ===")
    for i, full, hl in yellow_paras:
        print(f"★ [{i}] {full[:160]}")
        if hl.strip() != full.strip():
            # 只有部分 highlight 才印出 yellow snippet
            print(f"    黃字: {hl[:160]}")

    print(f"\n=== 刪除線段落（要砍）共 {len(strike_paras)} 處 ===")
    for i, full, st in strike_paras:
        print(f"✂ [{i}] full: {full[:160]}")
        if st.strip() != full.strip():
            print(f"    strike: {st[:160]}")

    print(f"\n=== 表格內標記 共 {len(tbl)} 處 ===")
    for loc, full, y, s in tbl[:50]:
        print(f"  [{loc}] {full[:120]}")
        if y.strip():
            print(f"     ★ {y[:120]}")
        if s.strip():
            print(f"     ✂ {s[:120]}")

    print(f"\n=== 嵌入圖片：{len(media)} 張 ===")
    if not media:
        print("  ⚠ 沒有實際嵌入圖、可能只有檔名提示。檢查文字中的圖檔名（如 '0303-shop.jpg'）。")
    for n in media:
        print(f"  📷 {n}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    main(sys.argv[1])
